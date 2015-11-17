
#   Author: Christopher Bull. 
#   Affiliation: Climate Change Research Centre and ARC Centre of Excellence for Climate System Science.
#                Level 4, Mathews Building
#                University of New South Wales
#                Sydney, NSW, Australia, 2052
#   Contact: z3457920@student.unsw.edu.au
#   www:     christopherbull.com.au
#   Date created: Fri, 13 Nov 2015 13:45:57
#   Machine created on: ccrc165
#

"""
Main script for MkPaper package
"""
from ._cblogger import _LogStart
import os
import contextlib as ctx
import subprocess

_lg=_LogStart().setup()

header=\
r"""
\title{Figures}
\author{
        This is an author name \
}
\date{\today}

\documentclass{scrartcl}

\usepackage{caption}  % Use for \captionof(*) command

% Set the page margins to accomodate space for captions and page numbers
\usepackage[left=0.9cm,right=0.9cm,bottom=1cm,footskip=1em,includefoot]{geometry}
% Use for 'landscape' environment to position landscape captions properly
\usepackage{pdflscape}

% Setup the new 'pagecommand*' option key-value
\usepackage{etoolbox}
\usepackage{pdfpages}
\makeatletter
\newcommand*{\AM@pagecommandstar}{}
\define@key{pdfpages}{pagecommand*}{\def\AM@pagecommandstar{#1}}
\patchcmd{\AM@output}{\begingroup\AM@pagecommand\endgroup}
{\ifthenelse{\boolean{AM@firstpage}}{\begingroup\AM@pagecommandstar\endgroup}{\begingroup\AM@pagecommand\endgroup}}{}{} % Patch to use new option
\patchcmd{\AM@split@optionsii}{\equal{pagecommand}{\AM@temp}\or}
{\equal{pagecommand}{\AM@temp}\or\equal{pagecommand*}{\AM@temp}\or}{}{}
\makeatother

\begin{document}
\maketitle

\begin{abstract}
These are the figures for the a paper\ldots
\end{abstract}

\section{Figure}
This is a descriptive text about the figures
"""

newfig=\
r"""
% Include a single PDF page
\includepdf[
    frame,
    scale=0.90,
    % Use new 'pagecommand*' to set caption on the first page, include page #
"""


def mkdir(p):
    """make directory of path that is passed"""
    try:
       os.makedirs(p)
       _lg.info("output folder: "+p+ " does not exist, we will make one.")
    except OSError as exc: # Python >2.5
       import errno
       if exc.errno == errno.EEXIST and os.path.isdir(p):
          pass
       else: raise


class LatexFigureDoc(object):
    """
    Class for creating LaTeX Figure Doc.

    Keep calling add_figure method as required.

    Call end_tex method to complete document

    Parameters
    ----------
    :latexfol: folder to put the latex and pdf stuff in. (Will be created if it doesn't exist.) Needs to be full path, not relative.
    :figdoctitle: name of figure document title (default is 'mkpaperfigure')

    Returns
    -------
    
    Notes
    -------

    Example
    --------
    >>> figobj=LatexFigureDoc(latexfol,figdoctitle)
    >>> figobj.add_figure(fpath,caption)
    >>> figobj.add_figure(fpath2,caption2)
    >>> figobj.add_figure(fpath3,caption3)
    >>> figobj.end_tex()
    >>> figobj.build_tex()

    #Or more concretely...
    >>> import mkpaper as mp
    >>> import glob
    >>> ifiles=sorted(glob.glob('/path/to/plots/' + '*.pdf' ))
    >>> assert(ifiles!=[]),"glob didn't find anything!"
    >>> figobj=mp.LatexFigureDoc('/home/nfs/username/examplemkpaper/','testname')
    >>> figobj.add_figure(ifiles[0],'fig1')
    >>> figobj.add_figure(ifiles[1],'fig2')
    >>> figobj.add_figure(ifiles[2],'fig3')
    >>> #etc
    >>> figobj.end_tex()
    >>> figobj.build_tex()
    """
    def __init__(self,latexfol,figdoctitle='mkpaperfigure'):
        _lg.info("Creating LaTeX file in: " + latexfol)
        self.latexfol=latexfol
        self.figdoctitle=figdoctitle
        self.outname=latexfol+figdoctitle+'.tex'
        mkdir(self.latexfol)
        with ctx.closing(open(self.outname,'w')) as handle:
            handle.write(header)

    def add_figure(self,filepath,caption):
        """function to add figure and caption to latex file.
        
        :filepath: this needs to be a pdf file!
        :caption:  caption for said figure
        :returns: 
        """
        _lg.debug("Adding Figure: " + filepath)
        with ctx.closing(open(self.outname,'a')) as handle:
            handle.write(newfig)
            newfig_intro=r"pagecommand*={\thispagestyle{plain}\null\vfill\captionof{figure}{"
            newfig_mid="}}]{"
            newfig_outro="}"
            handle.write(newfig_intro+caption+newfig_mid+filepath+newfig_outro+"\n") 
        
    def end_tex(self):
        """function to add figure and caption to latex file.
        
        :filepath: @todo
        :caption: @todo
        :returns: @todo
        """
        with ctx.closing(open(self.outname,'a')) as handle:
            handle.write("\end{document}"+"\n") 

    def build_tex(self):
        """function to call pdflatex to actually build the pdf file.
        
        :filepath: @todo
        :caption: @todo
        :returns: @todo
        """
        _lg.info("Building pdf file... (requires pdflatex)")

        #turned this off as it was crashing..
        #currentdir=os.getcwd()
        os.chdir(self.latexfol)
        pdfcommand=['pdflatex' , self.outname]

        FNULL = open(os.devnull, 'w')
        retcode = subprocess.call(pdfcommand, stdout=FNULL, stderr=subprocess.STDOUT)

        #os.chdir(currentdir)

        if os.path.isfile(self.outname[:-3]+'pdf'):
            _lg.info("MkPaper SUCCESS, check it out: "+self.outname[:-3]+'pdf')
        else:
            _lg.info("MkPaper FAIL")
            _lg.error("Something went wrong with pdflatex, it hasn't made a figure doc. We won't delete the LaTeX.")
            sys.exit("Something went wrong with pdflatex, it hasn't made a figure doc. We won't delete the LaTeX.")

class WordFigureDoc(object):
    """
    Class for creating Word Figure Doc.

    Keep calling add_figure method as required.

    Call end_doc method to complete document

    Parameters
    ----------
    :wordfol: folder to put the word doc in. (Will be created if it doesn't exist.) Needs to be full path, not relative.
    :figdoctitle: name of figure document title (default is 'mkpaperfigure')

    Returns
    -------
    
    Notes
    -------

    Example
    --------
    >>> figobj=WordFigureDoc(latexfol,figdoctitle)
    >>> figobj.add_figure(fpath,caption)
    >>> figobj.add_figure(fpath2,caption2,unicodeme=True)
    >>> figobj.end_doc()

    #Or more concretely...
    >>> import mkpaper as mp
    >>> import glob

    >>> #get a list of your figures in PNG format...
    >>> ifiles=sorted(glob.glob('/path/to/plots/' + '*.png' ))
    >>> assert(ifiles!=[]),"glob didn't find anything!"

    >>> figobj=mp.WordFigureDoc('/home/nfs/username/examplemkpaper/','testname')
    >>> figobj.add_figure(ifiles[0],'fig1')
    >>> figobj.add_figure(ifiles[1],'fig2')
    >>> figobj.add_figure(ifiles[2],'fig3')
    >>> #etc

    >>> #insert the tail of the doc
    >>> figobj.end_doc()
    """
    def __init__(self,wordfol,figdoctitle='mkpaperfigure'):
        _lg.info("Creating word doc file in: " + wordfol)
        self.wordfol=wordfol
        self.figdoctitle=figdoctitle
        self.outname=wordfol+figdoctitle+'.docx'
        mkdir(self.wordfol)
        from docx import Document

        self.document=Document()
        self.document.add_heading("Figure File", 0)
        self.document.add_paragraph("This doc is a figure container")

    def add_figure(self,filepath,caption,unicodeme=False):
        """function to add figure and caption to word doc file.
        
        :filepath: this needs to be a pdf file!
        :caption:  caption for said figure
        :unicodeme:  (optional) set to True if caption contains funky characters
        :returns: 

        Notes
        -------
        Insert this at the start of your python file to have funky strings..

            # coding=utf-8

        Discussion:
        http://stackoverflow.com/questions/3215168/how-to-get-charater-in-a-string-in-python
        """
        from docx.shared import Inches
        _lg.debug("Adding Figure: " + filepath)

        self.document.add_picture(filepath,width=Inches(6.0))
        if unicodeme:
            self.document.add_paragraph(unicode(caption,'utf-8'))
        else:
            self.document.add_paragraph(caption)
        self.document.add_page_break()
        
    def end_doc(self):
        self.document.save(self.outname)
        if os.path.isfile(self.outname):
            _lg.info("MkPaper SUCCESS, check it out: "+self.outname)
        else:
            _lg.info("MkPaper FAIL")
            _lg.error("Something went wrong with python-docx, it hasn't made a figure doc.")
            sys.exit("Something went wrong with python-docx, it hasn't made a figure doc.")


if __name__ == "__main__": 
    pass
